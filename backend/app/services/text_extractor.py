"""
Text Extraction Service
Handles extraction from PDF, Word, and Image files with OCR support
"""
import io
import re
from typing import Tuple, Optional
from pathlib import Path

import PyPDF2
import pdfplumber
from docx import Document
from PIL import Image
import pytesseract

from app.config import settings
from app.utils.logger import log_info, log_error, log_warning, log_exception


class TextExtractionError(Exception):
    """Custom exception for text extraction errors"""
    pass


class TextExtractor:
    """
    Robust text extractor supporting multiple formats
    Falls back to OCR when needed
    """
    
    def __init__(self):
        # Configure Tesseract if path is set
        if settings.TESSERACT_CMD:
            pytesseract.pytesseract.tesseract_cmd = settings.TESSERACT_CMD
        
        self.min_text_length = 50  # Minimum viable text length
    
    def extract(self, file_content: bytes, filename: str) -> str:
        """
        Main extraction method - routes to appropriate extractor
        
        Args:
            file_content: Binary file content
            filename: Original filename to determine type
            
        Returns:
            Extracted text
            
        Raises:
            TextExtractionError: If extraction fails
        """
        file_ext = Path(filename).suffix.lower()
        
        log_info(f"Starting text extraction for: {filename} (type: {file_ext})")
        
        try:
            if file_ext == '.pdf':
                text = self._extract_from_pdf(file_content, filename)
            elif file_ext in ['.docx', '.doc']:
                text = self._extract_from_word(file_content, filename)
            elif file_ext in ['.jpg', '.jpeg', '.png']:
                text = self._extract_from_image(file_content, filename)
            else:
                raise TextExtractionError(f"Unsupported file format: {file_ext}")
            
            # Validate extracted text
            if not text or len(text.strip()) < self.min_text_length:
                raise TextExtractionError(
                    f"Insufficient text extracted ({len(text)} chars). "
                    f"Minimum required: {self.min_text_length}"
                )
            
            cleaned_text = self._clean_text(text)
            log_info(f"Successfully extracted {len(cleaned_text)} characters from {filename}")
            
            return cleaned_text
            
        except Exception as e:
            log_exception(e, f"Failed to extract text from {filename}")
            raise TextExtractionError(f"Text extraction failed: {str(e)}")
    
    def _extract_from_pdf(self, file_content: bytes, filename: str) -> str:
        """
        Extract text from PDF using multiple methods
        1. Try pdfplumber (best for tables and layout)
        2. Fallback to PyPDF2 (faster, simpler)
        3. If minimal text, apply OCR (scanned PDFs)
        """
        text = ""
        
        # Method 1: pdfplumber (best quality)
        try:
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
                        log_debug(f"pdfplumber extracted {len(page_text)} chars from page {page_num}")
        except Exception as e:
            log_warning(f"pdfplumber failed for {filename}: {e}")
        
        # Method 2: PyPDF2 fallback
        if len(text.strip()) < self.min_text_length:
            log_info(f"pdfplumber yielded minimal text, trying PyPDF2")
            try:
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
            except Exception as e:
                log_warning(f"PyPDF2 failed for {filename}: {e}")
        
        # Method 3: OCR if still insufficient
        if len(text.strip()) < self.min_text_length and settings.OCR_ENABLED:
            log_info(f"Digital extraction yielded minimal text, attempting OCR")
            try:
                ocr_text = self._ocr_pdf(file_content, filename)
                if len(ocr_text) > len(text):
                    text = ocr_text
                    log_info(f"OCR successfully extracted text from scanned PDF")
            except Exception as e:
                log_error(f"OCR failed for {filename}: {e}")
        
        return text
    
    def _ocr_pdf(self, file_content: bytes, filename: str) -> str:
        """
        Apply OCR to scanned PDF
        Converts PDF pages to images and applies Tesseract
        """
        if not settings.OCR_ENABLED:
            raise TextExtractionError("OCR is disabled")
        
        try:
            from pdf2image import convert_from_bytes
        except ImportError:
            raise TextExtractionError(
                "pdf2image not installed. Install with: pip install pdf2image"
            )
        
        try:
            # Convert PDF to images
            images = convert_from_bytes(file_content, dpi=300)  # High DPI for better OCR
            text = ""
            
            for page_num, image in enumerate(images, 1):
                log_info(f"OCR processing page {page_num}/{len(images)}")
                
                # Preprocess image for better OCR
                image = self._preprocess_image(image)
                
                # Apply OCR
                page_text = pytesseract.image_to_string(image, lang='eng')
                text += page_text + "\n\n"
                
                log_debug(f"OCR extracted {len(page_text)} chars from page {page_num}")
            
            return text
            
        except Exception as e:
            raise TextExtractionError(f"PDF OCR failed: {str(e)}")
    
    def _extract_from_word(self, file_content: bytes, filename: str) -> str:
        """Extract text from Word document (.docx)"""
        try:
            doc = Document(io.BytesIO(file_content))
            text = ""
            
            # Extract from paragraphs
            for para in doc.paragraphs:
                text += para.text + "\n"
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                text += "\n"
            
            log_info(f"Extracted {len(text)} chars from Word document")
            return text
            
        except Exception as e:
            raise TextExtractionError(f"Word extraction failed: {str(e)}")
    
    def _extract_from_image(self, file_content: bytes, filename: str) -> str:
        """Extract text from image using OCR"""
        if not settings.OCR_ENABLED:
            raise TextExtractionError("OCR is disabled. Cannot process image files.")
        
        try:
            image = Image.open(io.BytesIO(file_content))
            
            # Preprocess for better OCR
            image = self._preprocess_image(image)
            
            # Apply OCR
            text = pytesseract.image_to_string(image, lang='eng')
            
            log_info(f"OCR extracted {len(text)} chars from image")
            return text
            
        except Exception as e:
            raise TextExtractionError(f"Image OCR failed: {str(e)}")
    
    def _preprocess_image(self, image: Image.Image) -> Image.Image:
        """
        Preprocess image for better OCR results
        - Convert to grayscale
        - Enhance contrast
        - Remove noise
        """
        try:
            # Convert to grayscale
            image = image.convert('L')
            
            # Optional: Apply threshold to make text more clear
            # This works well for scanned documents
            from PIL import ImageEnhance
            
            # Increase contrast
            enhancer = ImageEnhance.Contrast(image)
            image = enhancer.enhance(2.0)
            
            # Increase sharpness
            enhancer = ImageEnhance.Sharpness(image)
            image = enhancer.enhance(1.5)
            
            return image
            
        except Exception as e:
            log_warning(f"Image preprocessing failed: {e}")
            return image  # Return original if preprocessing fails
    
    def _clean_text(self, text: str) -> str:
        """
        Clean and normalize extracted text
        - Remove excessive whitespace
        - Fix common OCR errors
        - Normalize line breaks
        """
        if not text:
            return ""
        
        # Replace multiple whitespaces with single space
        text = re.sub(r'\s+', ' ', text)
        
        # Replace multiple newlines with double newline
        text = re.sub(r'\n+', '\n\n', text)
        
        # Remove non-printable characters but keep essential punctuation
        text = re.sub(r'[^\w\s.,;:()\-@#+*/\n]', '', text)
        
        # Fix common OCR errors
        text = self._fix_ocr_errors(text)
        
        return text.strip()
    
    def _fix_ocr_errors(self, text: str) -> str:
        """Fix common OCR recognition errors"""
        
        # Common OCR mistakes
        replacements = {
            r'\bl\b': 'I',  # Lowercase L mistaken for I
            r'\b0\b': 'O',  # Zero mistaken for O (context-dependent)
            r'rn': 'm',     # rn often mistaken for m
        }
        
        for pattern, replacement in replacements.items():
            text = re.sub(pattern, replacement, text)
        
        return text
    
    def extract_contact_info(self, text: str) -> dict:
        """
        Extract contact information from text
        Returns: dict with email, phone, linkedin
        """
        contact = {
            "email": None,
            "phone": None,
            "linkedin": None,
            "github": None
        }
        
        # Email regex
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        email_match = re.search(email_pattern, text)
        if email_match:
            contact["email"] = email_match.group()
        
        # Phone regex (various formats)
        phone_pattern = r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
        phone_match = re.search(phone_pattern, text)
        if phone_match:
            contact["phone"] = phone_match.group()
        
        # LinkedIn URL
        linkedin_pattern = r'(?:linkedin\.com/in/|linkedin\.com/pub/)[\w\-]+'
        linkedin_match = re.search(linkedin_pattern, text, re.IGNORECASE)
        if linkedin_match:
            contact["linkedin"] = linkedin_match.group()
        
        # GitHub URL
        github_pattern = r'github\.com/[\w\-]+'
        github_match = re.search(github_pattern, text, re.IGNORECASE)
        if github_match:
            contact["github"] = github_match.group()
        
        return contact


# Helper function for logging
def log_debug(message: str):
    """Placeholder for debug logging"""
    from app.utils.logger import log_debug as _log_debug
    _log_debug(message)